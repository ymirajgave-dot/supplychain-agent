import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from openai import OpenAI

st.set_page_config(
    page_title="Supply Chain Advanced Analytics — Project Document Generator",
    layout="centered"
)

# ---------- helpers ----------
def get_openai_key():
    try:
        return st.secrets.get("OPENAI_API_KEY", None)
    except Exception:
        return None

def read_any_file(upload):
    name = (upload.name or "").lower()
    data = upload.read()
    bio = BytesIO(data)
    if name.endswith(".csv"):
        try:
            df = pd.read_csv(bio)
            return "table", df
        except Exception:
            return "text", data.decode("utf-8", "ignore")
    if name.endswith(".xlsx"):
        try:
            df = pd.read_excel(bio)
            return "table", df
        except Exception:
            return "text", data.decode("utf-8", "ignore")
    if name.endswith(".txt"):
        return "text", data.decode("utf-8", "ignore")
    if name.endswith(".docx"):
        doc = Document(bio)
        txt = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return "text", txt
    return "text", data.decode("utf-8", "ignore")

def to_docx_bytes(text: str) -> bytes:
    doc = Document()
    for line in (text or "").splitlines():
        if line.strip().endswith(":") and not line.startswith(" "):
            # treat lines ending with ":" as headings
            doc.add_heading(line.strip(), level=2)
        elif line.strip().startswith("•") or line.strip().startswith("- "):
            p = doc.add_paragraph(line.strip().lstrip("•- "))
            try:
                p.style = "List Bullet"
            except Exception:
                pass
        else:
            doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

def polish_with_openai(notes, file_text, focus):
    key = get_openai_key()
    if not key:
        return None, "No OpenAI API key found"
    client = OpenAI(api_key=key)

    prompt = f"""
You are an expert in supply chain business strategy and technical data systems. 
Generate a clean, publication-ready organizational document for any supply chain initiative 
based on user-provided inputs (meeting transcripts, project documentation, discussions, and CSV data).

GOAL
Create a comprehensive Word-style document (no markdown, no decorative characters, no emojis, no hashtags, no asterisks) 
usable by both executives and technical teams. If a .docx writer is available, format headings and lists accordingly; 
otherwise output clean plain text that can be pasted into Word without edits.

INPUTS YOU WILL RECEIVE
• Free text: transcripts, notes, requirements, decisions
• Documents: policies, SOPs, scope statements (as text)
• Data tables (CSV excerpts pasted as text)

GAP DETECTION AND EXPERT SUGGESTIONS
• Proactively identify missing or weak elements (definitions, KPIs, targets, eligibility rules, governance, controls, UX specs, NFRs, risks).
• Add a clearly labeled “Expert Suggestions” subsection inside the relevant part of the document whenever the inputs are incomplete. 
Each suggestion must be practical, implementation-ready, and mapped to a measurable impact 
(time saved, cost avoided, revenue protected, working capital, service level).
• Where appropriate, offer alternative design options (Option A/B/C) with trade-offs (speed vs. accuracy, cost vs. performance, build vs. buy).

DATA INTERPRETATION RULES
Use only facts present in the inputs; do not invent numbers. If specifics are missing, state assumptions and data gaps and propose fixes.
Select KPIs appropriate to the project scope. Provide plain-text formulas.
When line-level data is provided, roll up to the correct KPI grain and document roll-up logic.
If inputs conflict, call out the conflict, list candidate systems of record, and recommend a governance step to resolve.

DOCUMENT STRUCTURE (produce in this exact order)

Title Page
Project title: {{project_name}}
Organization: {{company_name}}
Department: {{department_or_function}}
Document Owner: {{owner_role_or_name}}
Version: {{version}}
Date: {{today}}

Executive Summary (concise)
• Business problem and urgency
• Top three objectives and headline KPI(s)
• Expected business impact (time saved, cost avoided, revenue protected, working capital impact)

Table of Contents

Section A: Business Priorities
A.1 Project Overview and Context
A.2 Problem Statements
A.3 KPI Definitions and Targets
A.4 Business Case
A.5 Resolution Strategy
A.6 Milestones, Timelines, Ownership (RACI)
A.7 Anticipated Impact and Measurement

Section B: Execution Plan
B.1 Scope of Data and Grain
B.2 Source Systems and Integration Points
B.3 Target Data Model
B.4 Transformations and Business Rules
B.5 Pipeline Design
B.6 Analytics App / Power BI Design and UX
B.7 Advanced Analytics (optional)
B.8 Non-Functional Requirements
B.9 Risks, Dependencies, Mitigations
B.10 Phased Implementation Plan

Appendices
Glossary, Data dictionary, Assumptions, Sample calculations, Source Notes

FEEDBACK LOOP
End the document with a section titled Refinement Notes containing 5 concise questions tailored to the initiative.

STYLE AND FORMATTING
• Clear professional business English for executives and technical readers.
• Output must be Word-ready text with headings, subheadings, paragraphs, and simple lists.
• Do not use markdown syntax, code fences, emojis, hashtags, or decorative characters.

HALLUCINATION CONTROL & TRACEABILITY
• Use only provided facts; label examples as illustrative.
• If multiple inputs disagree, show both values, indicate uncertainty, and recommend a resolution path.

FINAL OUTPUT INSTRUCTIONS
• Produce only the document content (no meta commentary, no prompt text).
• If you can create a .docx file, do so; otherwise output clean text suitable for pasting into Word.
• Stop after the document content and the Refinement Notes.

=== INPUT START ===
Notes: {notes}

Uploaded Files (combined text or table samples):
{file_text or "(none)"}

Focus:
{focus or "(none)"}
=== INPUT END ===
"""

    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1800,
            temperature=0.2,
        )
        return resp.choices[0].message.content, None
    except Exception as e:
        return None, str(e)

# ---------- UI ----------
with st.sidebar:
    st.markdown("### 🔑 OpenAI Key")
    st.write("Status:", "✅ Found" if get_openai_key() else "❌ Missing")
    st.markdown("---")
    st.markdown("### ℹ️ About this Tool")
    st.info(
        "This tool converts raw supply chain notes and supporting files "
        "into a polished organizational document. "
        "It includes expert suggestions for gaps, KPIs, and implementation "
        "options, producing leadership-ready documentation in Word format."
    )

st.title("Supply Chain Advanced Analytics — Project Document Generator")

tab1, tab2 = st.tabs(["📂 Input", "📄 Output"])

# --- Tab 1: Inputs ---
with tab1:
    uploads = st.file_uploader(
        "Upload files (CSV/XLSX/TXT/DOCX) — multiple allowed",
        type=["csv", "xlsx", "txt", "docx"],
        accept_multiple_files=True
    )

    file_text = ""
    if uploads:
        parts = []
        for up in uploads:
            kind, content = read_any_file(up)
            if kind == "table":
                st.write(f"Table: {up.name}")
                st.dataframe(content.head())
                parts.append(content.head(20).to_csv(index=False))
            else:
                st.write(f"Text: {up.name}")
                preview = str(content)[:500]
                st.text_area(f"Preview: {up.name}", preview, height=120)
                parts.append(str(content))
        file_text = "\n\n".join(parts)

    with st.form("input_form"):
        notes = st.text_area(
            "Stakeholder notes",
            "Example: Stakeholder wants analytics MVP.",
            height=120
        )
        focus = st.text_area(
            "Additional focus (optional)",
            "",
            height=80
        )
        submitted = st.form_submit_button("Generate Project Document")

# --- Tab 2: Output ---
with tab2:
    if submitted:
        st.info("Generating with OpenAI…")
        polished, err = polish_with_openai(notes, file_text, focus)
        if polished:
            st.subheader("📌 Final Project Document")

            # Clean font style for readability
            st.markdown(
                """
                <style>
                .stMarkdown, .reportview-container {
                    font-family: "Helvetica Neue", Arial, sans-serif;
                    font-size: 16px;
                    line-height: 1.6;
                }
                h1, h2, h3 {
                    color: #2C3E50;
                    margin-top: 20px;
                }
                </style>
                """,
                unsafe_allow_html=True
            )

            st.text_area("Generated Document", polished, height=600)

            # Word download
            docx_bytes = to_docx_bytes(polished)
            st.download_button(
                "⬇️ Download as Word (.docx)",
                data=docx_bytes,
                file_name="project.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.error(err)
    else:
        st.warning("⚠️ Fill in inputs in the first tab and click *Generate Project Document*.")

